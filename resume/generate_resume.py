#!/usr/bin/env python3
"""
Generate EstebanMorenoResume.pdf  →  ../images/EstebanMorenoResume.pdf
        EstebanMorenoResume.docx  →  ./EstebanMorenoResume.docx

Usage:  python3 resume/generate_resume.py
"""

import os, sys
from pathlib import Path

RESUME_DIR = Path(__file__).parent
PDF_OUT    = RESUME_DIR.parent / "images" / "EstebanMorenoResume.pdf"
DOCX_OUT   = RESUME_DIR / "EstebanMorenoResume.docx"

# ── Content ────────────────────────────────────────────────────────────────────

NAME    = "Esteban Moreno"
TITLE   = "Senior DevOps Engineer"

CONTACT_LINE_1 = "morenoramirezesteban@gmail.com  |  linkedin.com/in/estebanmorenoramirez  |  github.com/estebanmorenoit"
CONTACT_LINE_2 = "blog.estebanmoreno.link  |  estebanmoreno.link  |  Remote / EU-based"

PROFILE = (
    "Senior DevOps Engineer with 3+ years of experience designing and operating cloud-native "
    "infrastructure on AWS. Specialising in Kubernetes platform engineering, GitOps-driven CI/CD, "
    "Terraform infrastructure-as-code, and production observability. Currently building and scaling "
    "EKS-based platforms at X-Tention. AWS Certified Solutions Architect \u2014 Associate. "
    "Published 108+ technical articles on AWS and DevOps. "
    "Open to remote or EU-based opportunities."
)

# Skills: (category, value) — rendered as "Category: value" inline
SKILLS = [
    ("Cloud",                  "AWS \u2014 EKS, Lambda, EC2, S3, DynamoDB, CloudFront, CloudWatch, IAM, VPC, Route\u00a053"),
    ("Infrastructure as Code", "Terraform, Terraform Cloud, HCL, Ansible"),
    ("Containers & Kubernetes","Kubernetes, Docker, Helm, ArgoCD, k3s, microk8s, Tanzu, Zarf, Docker Compose"),
    ("CI/CD",                  "GitLab CI/CD, GitHub Actions, ArgoCD (GitOps), Jenkins"),
    ("Observability",          "Prometheus, Grafana, Thanos, Loki, Promtail, Fluent Bit, Alertmanager, OpenTelemetry, OpenSearch"),
    ("Security & Identity",    "Keycloak (OIDC/SSO), AWS IAM, SAST, secret detection in pipelines"),
    ("OS & Scripting",         "Linux (RHEL, Ubuntu, Debian), Bash"),
]

CERTS = [
    ("AWS Certified Solutions Architect \u2014 Associate", "Amazon Web Services"),
    ("AWS Certified Cloud Practitioner",                   "Amazon Web Services"),
    ("LPI Linux Essentials",                               "Linux Professional Institute"),
    ("Terraform Beginner Bootcamp",                        "ExamPro"),
    ("Cisco Certified Network Associate (CCNA)",           "Cisco"),
    ("Cisco Cybersecurity Essentials",                     "Cisco"),
]

XTENTION_BULLETS = [
    "Architect and manage production EKS clusters using Helm and ArgoCD for GitOps-driven, "
    "declarative deployments across multiple environments.",
    "Design and maintain GitLab CI/CD pipelines with integrated SAST, secret detection, and "
    "compliance gates \u2014 reducing release risk while accelerating delivery velocity.",
    "Operate a full-stack observability platform (Prometheus, Thanos, Grafana, Loki, Promtail, "
    "Alertmanager) providing unified metrics, logs, and alerting across distributed workloads.",
    "Deploy and maintain Keycloak as centralised identity provider, enabling OIDC-based SSO for "
    "internal services and enforcing least-privilege access policies.",
    "Automate cloud infrastructure provisioning with Terraform and Terraform Cloud, eliminating "
    "manual steps and ensuring consistent, version-controlled environments.",
    "Evaluate and operate lightweight Kubernetes distributions (k3s, microk8s) and air-gapped "
    "tooling (Tanzu, Zarf) for varied deployment environments.",
    "Champion infrastructure-as-code and DevOps best practices across engineering teams: code "
    "review, runbook documentation, and on-call process improvements.",
]

AMAZON_BULLETS = [
    "Maintained high-availability IT infrastructure for Amazon fulfilment and sort centres "
    "operating 24/7, supporting thousands of associates per shift.",
    "Resolved hardware, network, and software incidents under strict SLA requirements in a "
    "high-throughput, time-critical logistics environment.",
    "Managed configuration of network devices, Linux servers, and end-user systems across "
    "large-scale warehouse facilities.",
    "Collaborated with AWS cloud and engineering teams on internal tooling and infrastructure "
    "automation initiatives.",
    "Delivered technical onboarding and training for new IT team members, reducing ramp-up time "
    "and improving team self-sufficiency.",
]

CHAPMANBDSP_BULLETS = [
    "Managed the International Helpdesk Team, delivering comprehensive IT support across "
    "6 sites in the UK and UAE.",
    "Led migration of the entire workforce to a remote IT environment, implementing VPN, VoIP, "
    "and Ivanti technologies \u2014 increasing remote work capability and productivity by 30%.",
    "Established IT security policies and procedures that achieved Cyber Essentials Plus "
    "certification, ensuring a robust and auditable security baseline.",
]

FLUID_BULLETS = [
    "Provided 1st and 2nd line IT support, resolving hardware, software, and network incidents "
    "across a design consultancy environment.",
    "Maintained workstations, managed user accounts, and supported day-to-day IT operations.",
]

EDUCATION = [
    (
        "BEng Telecommunications Engineering",
        "Miguel Hern\u00e1ndez University, Alicante, Spain",
        "2010 \u2013 2016",
        "Final project: Design, simulation and measurement of Voltage Controlled Oscillators on "
        "ISM Band \u2014 graded 9.5/10. Modules included Computer Networks, Analog & Digital "
        "Electronics, Programming Fundamentals, and Engineering Electromagnetics.",
    ),
]

PROJECTS = [
    (
        "Production-Grade Self-Hosted Home Lab on Raspberry Pi",
        "Docker Compose  |  Prometheus  |  Grafana  |  Loki  |  Tailscale  |  GitHub Actions",
        "Full self-hosted infrastructure stack on a Raspberry Pi: Docker Compose services, "
        "monitored via Prometheus/Grafana/Loki/Uptime Kuma, Caddy as reverse proxy, Tailscale for "
        "secure remote access, and automated deployments via GitHub Actions.",
        "github.com/estebanmorenoit/rpi-home-lab",
    ),
    (
        "Cloud Resume Challenge \u2014 Serverless Portfolio on AWS",
        "AWS  |  Terraform  |  Lambda  |  DynamoDB  |  CloudFront  |  GitHub Actions",
        "Serverless portfolio site on AWS: Lambda + DynamoDB visitor counter, CloudFront CDN, S3 "
        "static hosting, Route 53, full Terraform IaC, and a GitHub Actions CI/CD pipeline. "
        "Documented across a 6-part blog series.",
        "github.com/estebanmorenoit/estebanmoreno-portfolio",
    ),
]

WRITING = (
    "Published 108+ technical articles on Kubernetes, Terraform, AWS, CI/CD, and observability "
    "at blog.estebanmoreno.link \u2014 including series on the Cloud Resume Challenge and "
    "#90DaysOfDevOps. Completed the #90DaysOfDevOps challenge with daily hands-on labs across "
    "AWS, Linux, Docker, Kubernetes, Ansible, Jenkins, and Grafana."
)


# ══════════════════════════════════════════════════════════════════════════════
# PDF  (reportlab)
# ══════════════════════════════════════════════════════════════════════════════

def build_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, KeepTogether
    )

    W, H = A4
    ACCENT  = colors.HexColor("#047857")
    DARK    = colors.HexColor("#111827")
    MID     = colors.HexColor("#374151")
    MUTED   = colors.HexColor("#6b7280")
    RULE    = colors.HexColor("#d1d5db")
    WHITE   = colors.white

    # ── Paragraph styles ──────────────────────────────────────────────────
    def ps(name, **kw):
        return ParagraphStyle(name, **kw)

    S_NAME     = ps("name",    fontName="Helvetica-Bold",    fontSize=24, textColor=DARK,   leading=28, spaceAfter=3)
    S_TITLE    = ps("title",   fontName="Helvetica",         fontSize=12, textColor=ACCENT, leading=16, spaceAfter=3)
    S_CONTACT  = ps("contact", fontName="Helvetica",         fontSize=8.5,textColor=MUTED,  leading=13, spaceAfter=2)
    S_SEC      = ps("sec",     fontName="Helvetica-Bold",    fontSize=9,  textColor=ACCENT, leading=12,
                               spaceBefore=10, spaceAfter=4,
                               textTransform="uppercase", letterSpacing=0.8)
    S_BODY     = ps("body",    fontName="Helvetica",         fontSize=9,  textColor=MID,    leading=14,
                               spaceAfter=4, alignment=TA_JUSTIFY)
    S_SKILL    = ps("skill",   fontName="Helvetica",         fontSize=8.5,textColor=MID,    leading=13, spaceAfter=3)
    S_JOB      = ps("job",     fontName="Helvetica-Bold",    fontSize=10, textColor=DARK,   leading=14,
                               spaceBefore=8, spaceAfter=1)
    S_META     = ps("meta",    fontName="Helvetica-Oblique", fontSize=8.5,textColor=MUTED,  leading=12, spaceAfter=4)
    # ── Hanging indent bullet — "• " is part of the text string ─────────
    S_BUL      = ps("bul",     fontName="Helvetica",         fontSize=8.5,textColor=MID,    leading=13,
                               leftIndent=10, firstLineIndent=-10,
                               spaceAfter=2)
    S_PROJ_T   = ps("projt",   fontName="Helvetica-Bold",    fontSize=9,  textColor=DARK,   leading=13,
                               spaceBefore=7, spaceAfter=1)
    S_PROJ_S   = ps("projs",   fontName="Helvetica-Oblique", fontSize=8,  textColor=ACCENT, leading=12, spaceAfter=2)
    S_PROJ_L   = ps("projl",   fontName="Helvetica",         fontSize=8,  textColor=MUTED,  leading=12, spaceAfter=2)
    S_CERT     = ps("cert",    fontName="Helvetica",         fontSize=8.5,textColor=MID,    leading=13, spaceAfter=2)

    def hr(color=RULE, thickness=0.6, before=2, after=4):
        return HRFlowable(width="100%", thickness=thickness, color=color,
                          spaceBefore=before, spaceAfter=after)

    def section(title):
        return [Paragraph(title, S_SEC), hr()]

    def bul(text):
        """Hanging-indent bullet — no XML <bullet> tag, so ATS parses cleanly."""
        return Paragraph("\u2022\u00a0" + text, S_BUL)

    # ── Document ──────────────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=18*mm, rightMargin=18*mm,
        topMargin=16*mm,  bottomMargin=16*mm,
    )

    story = []

    # Header ──────────────────────────────────────────────────────────────
    story.append(Paragraph(NAME, S_NAME))
    story.append(Paragraph(TITLE, S_TITLE))
    story.append(hr(color=colors.HexColor("#047857"), thickness=1.5, before=0, after=5))
    story.append(Paragraph(CONTACT_LINE_1, S_CONTACT))
    story.append(Paragraph(CONTACT_LINE_2, S_CONTACT))
    story.append(Spacer(1, 6))

    # Professional Profile ────────────────────────────────────────────────
    story += section("Professional Profile")
    story.append(Paragraph(PROFILE, S_BODY))

    # Core Skills ─────────────────────────────────────────────────────────
    story += section("Core Skills")
    for label, value in SKILLS:
        story.append(Paragraph(
            f"<b>{label}:</b>\u00a0\u00a0{value}",
            S_SKILL
        ))

    # Certifications ──────────────────────────────────────────────────────
    story += section("Certifications")
    for cert_name, issuer in CERTS:
        story.append(Paragraph(
            f"\u2022\u00a0<b>{cert_name}</b> \u2014 <i>{issuer}</i>",
            S_BUL
        ))

    # Work Experience ─────────────────────────────────────────────────────
    story += section("Work Experience")

    xtention_block = KeepTogether([
        Paragraph("Senior DevOps Engineer", S_JOB),
        Paragraph("X-Tention  \u00b7  Feb 2024 \u2013 Present", S_META),
        bul(XTENTION_BULLETS[0]),
        bul(XTENTION_BULLETS[1]),
        bul(XTENTION_BULLETS[2]),
    ])
    xtention_rest = [bul(b) for b in XTENTION_BULLETS[3:]]

    amazon_block = KeepTogether([
        Paragraph("IT Engineer", S_JOB),
        Paragraph("Amazon  \u00b7  Mar 2022 \u2013 Feb 2024", S_META),
        bul(AMAZON_BULLETS[0]),
        bul(AMAZON_BULLETS[1]),
    ])
    amazon_rest = [bul(b) for b in AMAZON_BULLETS[2:]]

    story.append(xtention_block)
    story += xtention_rest
    story.append(amazon_block)
    story += amazon_rest

    chapman_block = KeepTogether([
        Paragraph("Senior IT Engineer", S_JOB),
        Paragraph("chapmanbdsp  \u00b7  Mar 2018 \u2013 Mar 2022", S_META),
        bul(CHAPMANBDSP_BULLETS[0]),
        bul(CHAPMANBDSP_BULLETS[1]),
    ])
    story.append(chapman_block)
    story.append(bul(CHAPMANBDSP_BULLETS[2]))

    fluid_block = KeepTogether([
        Paragraph("IT Support Engineer", S_JOB),
        Paragraph("Fluid Design Ltd  \u00b7  Aug 2017 \u2013 Mar 2018", S_META),
        bul(FLUID_BULLETS[0]),
        bul(FLUID_BULLETS[1]),
    ])
    story.append(fluid_block)

    # Education ───────────────────────────────────────────────────────────
    story += section("Education")
    for degree, institution, dates, detail in EDUCATION:
        story.append(KeepTogether([
            Paragraph(degree, S_JOB),
            Paragraph(f"{institution}  \u00b7  {dates}", S_META),
            Paragraph(detail, S_BODY),
        ]))

    # Selected Projects ───────────────────────────────────────────────────
    story += section("Selected Projects")
    for ptitle, pstack, pdesc, plink in PROJECTS:
        story.append(KeepTogether([
            Paragraph(ptitle, S_PROJ_T),
            Paragraph(pstack, S_PROJ_S),
            Paragraph(pdesc,  S_BODY),
            Paragraph(plink,  S_PROJ_L),
        ]))

    # Writing & Community ─────────────────────────────────────────────────
    story += section("Writing & Community")
    story.append(Paragraph(WRITING, S_BODY))

    doc.build(story)
    print(f"PDF  \u2192 {PDF_OUT}")


# ══════════════════════════════════════════════════════════════════════════════
# DOCX  (python-docx)
# ══════════════════════════════════════════════════════════════════════════════

def build_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ACCENT_RGB = RGBColor(0x04, 0x78, 0x57)
    DARK_RGB   = RGBColor(0x11, 0x18, 0x27)
    MID_RGB    = RGBColor(0x37, 0x41, 0x51)
    MUTED_RGB  = RGBColor(0x6b, 0x72, 0x80)

    doc = Document()
    for sec in doc.sections:
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)
        sec.top_margin    = Cm(1.6)
        sec.bottom_margin = Cm(1.6)

    def add_run(para, text, bold=False, italic=False, size=10, color=None, font_name="Arial"):
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = font_name
        if color:
            run.font.color.rgb = color
        return run

    def para(space_before=0, space_after=2):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        return p

    def add_hr(p=None):
        if p is None:
            p = para(space_before=0, space_after=3)
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "D1D5DB")
        pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
        return p

    def section_heading(title):
        p = para(space_before=8, space_after=1)
        add_run(p, title.upper(), bold=True, size=9, color=ACCENT_RGB)
        add_hr()

    def bullet_item(text, size=8.5):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        add_run(p, text, size=size, color=MID_RGB)

    # Header
    p = para(space_after=2)
    add_run(p, NAME, bold=True, size=22, color=DARK_RGB)

    p = para(space_after=2)
    add_run(p, TITLE, size=12, color=ACCENT_RGB)

    add_hr()

    p = para(space_after=1)
    add_run(p, CONTACT_LINE_1, size=8.5, color=MUTED_RGB)

    p = para(space_after=5)
    add_run(p, CONTACT_LINE_2, size=8.5, color=MUTED_RGB)

    # Profile
    section_heading("Professional Profile")
    p = para(space_after=4)
    add_run(p, PROFILE, size=9, color=MID_RGB)

    # Skills
    section_heading("Core Skills")
    for label, value in SKILLS:
        p = para(space_after=2)
        add_run(p, label + ":  ", bold=True, size=8.5, color=DARK_RGB)
        add_run(p, value, size=8.5, color=MID_RGB)

    # Certs
    section_heading("Certifications")
    for cert_name, issuer in CERTS:
        bullet_item(cert_name + " \u2014 " + issuer)

    # Experience
    section_heading("Work Experience")

    p = para(space_before=4, space_after=1)
    add_run(p, "Senior DevOps Engineer", bold=True, size=10, color=DARK_RGB)

    p = para(space_after=3)
    add_run(p, "X-Tention  \u00b7  Feb 2024 \u2013 Present", italic=True, size=8.5, color=MUTED_RGB)

    for b in XTENTION_BULLETS:
        bullet_item(b)

    p = para(space_before=6, space_after=1)
    add_run(p, "IT Engineer", bold=True, size=10, color=DARK_RGB)

    p = para(space_after=3)
    add_run(p, "Amazon  \u00b7  Mar 2022 \u2013 Feb 2024", italic=True, size=8.5, color=MUTED_RGB)

    for b in AMAZON_BULLETS:
        bullet_item(b)

    p = para(space_before=6, space_after=1)
    add_run(p, "Senior IT Engineer", bold=True, size=10, color=DARK_RGB)

    p = para(space_after=3)
    add_run(p, "chapmanbdsp  \u00b7  Mar 2018 \u2013 Mar 2022", italic=True, size=8.5, color=MUTED_RGB)

    for b in CHAPMANBDSP_BULLETS:
        bullet_item(b)

    p = para(space_before=6, space_after=1)
    add_run(p, "IT Support Engineer", bold=True, size=10, color=DARK_RGB)

    p = para(space_after=3)
    add_run(p, "Fluid Design Ltd  \u00b7  Aug 2017 \u2013 Mar 2018", italic=True, size=8.5, color=MUTED_RGB)

    for b in FLUID_BULLETS:
        bullet_item(b)

    # Education
    section_heading("Education")
    for degree, institution, dates, detail in EDUCATION:
        p = para(space_before=4, space_after=1)
        add_run(p, degree, bold=True, size=10, color=DARK_RGB)

        p = para(space_after=3)
        add_run(p, f"{institution}  \u00b7  {dates}", italic=True, size=8.5, color=MUTED_RGB)

        p = para(space_after=4)
        add_run(p, detail, size=8.5, color=MID_RGB)

    # Projects
    section_heading("Selected Projects")
    for ptitle, pstack, pdesc, plink in PROJECTS:
        p = para(space_before=4, space_after=1)
        add_run(p, ptitle, bold=True, size=9, color=DARK_RGB)

        p = para(space_after=1)
        add_run(p, pstack, italic=True, size=8, color=ACCENT_RGB)

        p = para(space_after=1)
        add_run(p, pdesc, size=8.5, color=MID_RGB)

        p = para(space_after=3)
        add_run(p, plink, size=8, color=MUTED_RGB)

    # Writing
    section_heading("Writing & Community")
    p = para(space_after=4)
    add_run(p, WRITING, size=9, color=MID_RGB)

    doc.save(str(DOCX_OUT))
    print(f"DOCX \u2192 {DOCX_OUT}")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    RESUME_DIR.mkdir(parents=True, exist_ok=True)

    errors = []
    try:
        build_pdf()
    except Exception as e:
        errors.append(f"PDF: {e}")
        import traceback; traceback.print_exc()

    try:
        build_docx()
    except Exception as e:
        errors.append(f"DOCX: {e}")
        import traceback; traceback.print_exc()

    if errors:
        print("ERRORS:", errors); sys.exit(1)
    print("Done.")
